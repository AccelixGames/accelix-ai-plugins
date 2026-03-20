"""
Unified file reader for PDF, DOCX, XLSX, HWPX, HWP formats.
Returns extracted text as a UTF-8 string.
"""

import sys
import os
import zipfile
import xml.etree.ElementTree as ET

# Windows cp949 encoding defense
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')
if hasattr(sys.stderr, 'reconfigure'):
    sys.stderr.reconfigure(encoding='utf-8')


def read_file(path: str) -> str:
    """
    Read a file and return its text content.

    Supports: .pdf, .docx, .xlsx, .xls, .hwpx, .hwp

    Args:
        path: Absolute or relative path to the file.

    Returns:
        Extracted text as a UTF-8 string.

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the file format is unsupported.
        RuntimeError: If extraction fails (includes actionable message).
    """
    path = os.path.normpath(path)

    if not os.path.exists(path):
        raise FileNotFoundError(f"File not found: {path}")

    ext = os.path.splitext(path)[1].lower()

    readers = {
        '.pdf':  _read_pdf,
        '.docx': _read_docx,
        '.xlsx': _read_excel,
        '.xls':  _read_excel,
        '.hwpx': _read_hwpx,
        '.hwp':  _read_hwp,
    }

    if ext not in readers:
        raise ValueError(
            f"Unsupported file format: '{ext}'. "
            f"Supported formats: {', '.join(readers)}"
        )

    try:
        return readers[ext](path)
    except (ImportError, FileNotFoundError, ValueError):
        raise
    except Exception as e:
        raise RuntimeError(
            f"Failed to read '{os.path.basename(path)}': {e}. "
            f"Check reader patterns for format '{ext}'."
        ) from e


# ---------------------------------------------------------------------------
# Internal readers
# ---------------------------------------------------------------------------

def _read_pdf(path: str) -> str:
    """Extract text from a PDF using PyMuPDF (fitz) or pdfplumber."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(path)
        pages = [page.get_text() for page in doc]
        doc.close()
        return "\n".join(pages)
    except ImportError:
        pass

    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
        return "\n".join(pages)
    except ImportError:
        raise ImportError(
            "No PDF library found. Install one:\n"
            "  pip install pymupdf\n"
            "  pip install pdfplumber"
        )


def _read_docx(path: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        import docx
    except ImportError:
        raise ImportError(
            "python-docx is required for .docx files.\n"
            "  pip install python-docx"
        )

    doc = docx.Document(path)
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)

    return "\n".join(parts)


def _read_excel(path: str) -> str:
    """Extract text from an XLSX file using openpyxl."""
    ext = os.path.splitext(path)[1].lower()

    if ext == '.xls':
        raise ImportError(
            ".xls format requires xlrd.\n"
            "  pip install xlrd\n"
            "Note: xlrd 2.x only supports .xls (not .xlsx)."
        )

    try:
        import openpyxl
    except ImportError:
        raise ImportError(
            "openpyxl is required for .xlsx files.\n"
            "  pip install openpyxl"
        )

    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    parts = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"[Sheet: {sheet_name}]")
        for row in ws.iter_rows(values_only=True):
            row_text = "\t".join(
                str(cell) if cell is not None else "" for cell in row
            )
            if row_text.strip():
                parts.append(row_text)

    wb.close()
    return "\n".join(parts)


def _read_hwpx(path: str) -> str:
    """Extract text from an HWPX file (ZIP + XML, no external deps)."""
    HWPX_NS = "http://www.hancom.co.kr/hwpml/2012/paragraph"
    TEXT_TAGS = {
        f"{{{HWPX_NS}}}t",
        f"{{{HWPX_NS}}}T",
    }

    texts = []

    with zipfile.ZipFile(path, 'r') as zf:
        section_files = sorted(
            name for name in zf.namelist()
            if name.startswith("Contents/section") and name.endswith(".xml")
        )

        if not section_files:
            # Fallback: search any XML inside Contents/
            section_files = sorted(
                name for name in zf.namelist()
                if name.startswith("Contents/") and name.endswith(".xml")
            )

        for section_name in section_files:
            with zf.open(section_name) as f:
                tree = ET.parse(f)
                root = tree.getroot()
                for elem in root.iter():
                    if elem.tag in TEXT_TAGS and elem.text:
                        texts.append(elem.text)

    return "\n".join(texts)


def _read_hwp(path: str) -> str:
    """Extract text from a legacy HWP (OLE compound) file using olefile."""
    import struct
    import zlib

    try:
        import olefile
    except ImportError:
        raise ImportError(
            "olefile is required for .hwp files.\n"
            "  pip install olefile"
        )

    HWP_PARA_TEXT_TAG = 67

    with olefile.OleFileIO(path) as ole:
        # Read FileHeader to determine compression flag
        header_data = ole.openstream('FileHeader').read()
        # Compression flag is bit 1 of the flags dword at offset 36
        flags = struct.unpack_from('<I', header_data, 36)[0]
        compressed = bool(flags & 0x1)

        # Read BodyText/Section0
        if not ole.exists('BodyText/Section0'):
            return ""

        raw = ole.openstream('BodyText/Section0').read()

        if compressed:
            raw = zlib.decompress(raw, -15)

        texts = []
        offset = 0
        while offset + 4 <= len(raw):
            tag_header = struct.unpack_from('<I', raw, offset)[0]
            offset += 4

            tag_id   = tag_header & 0x3FF
            _level   = (tag_header >> 10) & 0x3FF
            size     = (tag_header >> 20) & 0xFFF

            if size == 0xFFF:
                if offset + 4 > len(raw):
                    break
                size = struct.unpack_from('<I', raw, offset)[0]
                offset += 4

            data = raw[offset:offset + size]
            offset += size

            if tag_id == HWP_PARA_TEXT_TAG:
                try:
                    text = data.decode('utf-16-le', errors='ignore')
                    # Filter non-printable control characters
                    text = "".join(
                        ch for ch in text
                        if ord(ch) >= 32 or ch in ('\n', '\t')
                    )
                    if text.strip():
                        texts.append(text)
                except Exception:
                    continue

    return "\n".join(texts)


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python read_file.py <file_path>", file=sys.stderr)
        sys.exit(1)
    try:
        result = read_file(sys.argv[1])
        print(result)
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)
